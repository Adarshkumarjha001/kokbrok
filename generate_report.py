from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create document
doc = Document()

# Title Page
title = doc.add_heading('Kokborok Language NLP Model', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('Part-of-Speech Tagging System')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(16)

doc.add_paragraph('\n\n')
doc.add_paragraph('Capstone Project Report').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('\n\n\n')

# Add project details
details = doc.add_paragraph()
details.alignment = WD_ALIGN_PARAGRAPH.CENTER
details.add_run('Submitted by: [Your Name]\n')
details.add_run('Roll No: [Your Roll Number]\n')
details.add_run('Department: Computer Science & Engineering\n')
details.add_run('Academic Year: 2024-2025\n')

doc.add_page_break()

# Table of Contents
doc.add_heading('Table of Contents', 1)
toc_items = [
    '1. Introduction',
    '   1.1 Objective of the Project',
    '   1.2 Description of the Project',
    '   1.3 Scope of the Project',
    '2. System Description',
    '   2.1 Customer/User Profiles',
    '   2.2 Assumptions and Dependencies',
    '   2.3 Functional Requirements',
    '   2.4 Non-Functional Requirements',
    '3. Design',
    '   3.1 System Design',
    '   3.2 Database Design',
    '4. Scheduling and Estimates'
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number' if not item.startswith('   ') else 'List Bullet')

doc.add_page_break()

# 1. INTRODUCTION
doc.add_heading('1. Introduction', 1)

# 1.1 Objective
doc.add_heading('1.1 Objective of the Project', 2)
obj_text = """The primary objective of this project is to develop an advanced Natural Language Processing (NLP) system for the Kokborok language, specifically focusing on Part-of-Speech (POS) tagging. The key objectives include:

• Preserve and promote the Kokborok language through digital technology
• Create an accurate POS tagging model using state-of-the-art transformer architecture
• Develop an accessible web interface for language analysis
• Provide educational resources about Kokborok language and its cultural significance
• Enable future NLP applications for this low-resource indigenous language
• Bridge the digital divide for Kokborok-speaking communities"""
doc.add_paragraph(obj_text)

# 1.2 Description
doc.add_heading('1.2 Description of the Project', 2)
desc_text = """This project implements a comprehensive NLP solution for Kokborok, a Tibeto-Burman language spoken by approximately 1 million people in Tripura, India. The system consists of:

1. Machine Learning Model: An XLM-RoBERTa transformer model fine-tuned for token classification, capable of identifying 17 different parts of speech in Kokborok text.

2. Web Application: A multi-page responsive website featuring:
   - Interactive demo for real-time POS tagging
   - Comprehensive language history documentation
   - Educational content about the importance of indigenous language preservation
   - Technical documentation of the NLP model

3. Deployment Infrastructure: The system is designed for cloud deployment using Hugging Face Spaces and GitHub Pages, ensuring accessibility and scalability.

The model leverages transfer learning from multilingual pre-training, making it effective despite Kokborok being a low-resource language."""
doc.add_paragraph(desc_text)

# 1.3 Scope
doc.add_heading('1.3 Scope of the Project', 2)
scope_text = """The scope of this project encompasses:

In Scope:
• Part-of-Speech tagging for Kokborok text with 17 grammatical categories
• Web-based user interface for text analysis
• Educational content about Kokborok language and culture
• Model deployment on cloud infrastructure
• Support for both Bengali and Roman scripts
• Real-time analysis with confidence scores
• Responsive design for mobile and desktop access

Out of Scope:
• Machine translation between languages
• Speech recognition or text-to-speech
• Named Entity Recognition (NER)
• Sentiment analysis
• Grammar correction
• Language learning applications

Future Enhancements:
• Expansion to other Tibeto-Burman languages
• Integration with translation systems
• Mobile application development
• Voice interface capabilities"""
doc.add_paragraph(scope_text)

doc.add_page_break()

# 2. SYSTEM DESCRIPTION
doc.add_heading('2. System Description', 1)

# 2.1 User Profiles
doc.add_heading('2.1 Customer/User Profiles', 2)
users_text = """The system is designed for multiple user categories:

1. Linguists and Researchers
   - Need: Analyze large corpora of Kokborok text
   - Usage: Research language patterns, document dialects
   - Technical Level: High

2. Educators and Students
   - Need: Learn and teach Kokborok grammar
   - Usage: Educational demonstrations, homework assistance
   - Technical Level: Medium

3. Native Speakers
   - Need: Digital tools in their native language
   - Usage: Text analysis, language validation
   - Technical Level: Low to Medium

4. Government Officials
   - Need: Process Kokborok documents
   - Usage: Administrative text processing
   - Technical Level: Medium

5. Technology Developers
   - Need: Build applications for Kokborok
   - Usage: API integration, further development
   - Technical Level: High"""
doc.add_paragraph(users_text)

# 2.2 Assumptions
doc.add_heading('2.2 Assumptions and Dependencies', 2)
assumptions_text = """Assumptions:
• Users have basic internet connectivity
• Input text is in Kokborok language
• Users understand the concept of parts of speech
• Text input is reasonably well-formed

Dependencies:
• Hugging Face Transformers library (v4.46+)
• PyTorch deep learning framework (v2.6+)
• Gradio web framework for UI
• Pre-trained XLM-RoBERTa model
• Internet connection for model loading
• Modern web browser (Chrome, Firefox, Safari, Edge)
• Python 3.11+ runtime environment"""
doc.add_paragraph(assumptions_text)

# 2.3 Functional Requirements
doc.add_heading('2.3 Functional Requirements', 2)
func_req = """FR1: Text Input Processing
   - System shall accept Kokborok text input up to 512 tokens
   - System shall handle both Bengali and Roman scripts
   - System shall validate input before processing

FR2: POS Tagging
   - System shall classify each token into one of 17 POS categories
   - System shall provide confidence scores for each prediction
   - System shall handle subword tokenization correctly

FR3: Result Display
   - System shall display tagged tokens with color coding
   - System shall show POS labels and confidence scores
   - System shall provide a legend explaining all tags

FR4: User Interface
   - System shall provide a responsive web interface
   - System shall support multiple pages (Home, History, Importance, Technology)
   - System shall include example sentences for testing

FR5: Educational Content
   - System shall display Kokborok language history
   - System shall explain the importance of language preservation
   - System shall document technical architecture

FR6: Performance
   - System shall process text within 3 seconds
   - System shall handle concurrent users
   - System shall maintain 95%+ uptime"""
doc.add_paragraph(func_req)

# 2.4 Non-Functional Requirements
doc.add_heading('2.4 Non-Functional Requirements', 2)
nonfunc_req = """NFR1: Usability
   - Interface shall be intuitive for non-technical users
   - Response time shall be under 3 seconds
   - Error messages shall be clear and helpful

NFR2: Reliability
   - System uptime shall be 99%+
   - Model accuracy shall be 85%+ on test data
   - System shall handle errors gracefully

NFR3: Performance
   - Page load time shall be under 2 seconds
   - System shall support 100+ concurrent users
   - Model inference shall complete within 1 second

NFR4: Scalability
   - Architecture shall support horizontal scaling
   - System shall handle increased traffic
   - Storage shall accommodate growing datasets

NFR5: Security
   - User input shall be sanitized
   - API endpoints shall be protected
   - No sensitive data shall be stored

NFR6: Maintainability
   - Code shall be well-documented
   - Architecture shall be modular
   - Updates shall be deployable without downtime

NFR7: Accessibility
   - Interface shall be mobile-responsive
   - Color contrast shall meet WCAG standards
   - Text shall be readable at various sizes"""
doc.add_paragraph(nonfunc_req)

doc.add_page_break()

# 3. DESIGN
doc.add_heading('3. Design', 1)

# 3.1 System Design
doc.add_heading('3.1 System Design', 2)
design_text = """The system follows a three-tier architecture:

1. Presentation Layer (Frontend)
   - HTML5/CSS3/JavaScript web interface
   - Responsive design using CSS Grid and Flexbox
   - Multiple pages: Home, History, Importance, Technology
   - Interactive demo with real-time feedback

2. Application Layer (Backend)
   - Gradio framework for API and UI
   - Python-based request handling
   - Model inference pipeline
   - Error handling and validation

3. Model Layer
   - XLM-RoBERTa transformer (125M parameters)
   - 12 transformer layers, 768 hidden dimensions
   - SentencePiece tokenizer
   - Token classification head with 17 output classes

System Components:
• Web Server: Gradio/Flask serving HTTP requests
• ML Model: Hugging Face Transformers pipeline
• Frontend: Static HTML/CSS/JS files
• API: RESTful endpoints for text analysis
• Storage: Model weights hosted on Hugging Face Hub"""
doc.add_paragraph(design_text)

# 3.1.1 Data Flow
doc.add_heading('3.1.1 Data Flow Diagram', 3)
dfd_text = """Level 0 DFD (Context Diagram):
User → [Kokborok NLP System] → Tagged Text Output

Level 1 DFD:
1. User Input → Text Validation → Validated Text
2. Validated Text → Tokenization → Token Sequence
3. Token Sequence → Model Inference → POS Predictions
4. POS Predictions → Result Formatting → Formatted Output
5. Formatted Output → Display → User Interface

Data Stores:
• Model Weights (Read-only)
• Configuration Files (Read-only)
• Temporary Session Data (In-memory)"""
doc.add_paragraph(dfd_text)

# 3.2 Database Design
doc.add_heading('3.2 Database Design', 2)
db_text = """This system primarily operates as a stateless application with minimal data persistence requirements.

Data Structures:

1. Model Configuration (JSON)
   - Architecture parameters
   - Label mappings (id2label, label2id)
   - Tokenizer settings
   - Model hyperparameters

2. POS Tag Definitions (In-memory)
   - Tag code (e.g., "NOUN", "VERB")
   - Full name (e.g., "Noun", "Verb")
   - Description
   - Color code for visualization
   - Example usage

3. Session Data (Temporary)
   - Input text
   - Tokenized sequence
   - Model predictions
   - Confidence scores
   - Formatted results

No persistent database is required as:
• Model weights are loaded from Hugging Face Hub
• No user authentication or data storage needed
• All processing is stateless and real-time
• Configuration is file-based

Future Database Considerations:
If user accounts or history tracking is added:
• User table (id, username, email, created_at)
• Analysis_history table (id, user_id, input_text, results, timestamp)
• Feedback table (id, user_id, analysis_id, rating, comments)"""
doc.add_paragraph(db_text)

doc.add_page_break()

# 4. SCHEDULING AND ESTIMATES
doc.add_heading('4. Scheduling and Estimates', 1)

# Add table for project schedule
table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Phase'
hdr_cells[1].text = 'Tasks'
hdr_cells[2].text = 'Duration'
hdr_cells[3].text = 'Status'

schedule_data = [
    ('Phase 1: Research & Planning', 'Literature review, requirement analysis, technology selection', '2 weeks', 'Completed'),
    ('Phase 2: Data Collection', 'Gather Kokborok text corpus, annotate POS tags', '3 weeks', 'Completed'),
    ('Phase 3: Model Development', 'Fine-tune XLM-RoBERTa, evaluate performance', '4 weeks', 'Completed'),
    ('Phase 4: Frontend Development', 'Design UI, implement HTML/CSS/JS pages', '2 weeks', 'Completed'),
    ('Phase 5: Backend Integration', 'Develop API, integrate model with frontend', '2 weeks', 'Completed'),
    ('Phase 6: Testing', 'Unit testing, integration testing, user testing', '2 weeks', 'In Progress'),
    ('Phase 7: Deployment', 'Deploy to Hugging Face Spaces, documentation', '1 week', 'Planned'),
    ('Phase 8: Maintenance', 'Bug fixes, performance optimization', 'Ongoing', 'Planned')
]

for phase, tasks, duration, status in schedule_data:
    row_cells = table.add_row().cells
    row_cells[0].text = phase
    row_cells[1].text = tasks
    row_cells[2].text = duration
    row_cells[3].text = status

doc.add_paragraph('\n')

estimates_text = """Resource Estimates:

Human Resources:
• 1 ML Engineer: Model development and training
• 1 Frontend Developer: UI/UX design and implementation
• 1 Backend Developer: API and integration
• 1 Linguist: Data annotation and validation
• Total: 4 team members

Time Estimates:
• Total Project Duration: 16 weeks
• Development Phase: 11 weeks
• Testing Phase: 2 weeks
• Deployment Phase: 1 week
• Buffer Time: 2 weeks

Cost Estimates:
• Cloud Computing (Training): $200
• Hugging Face Spaces Hosting: Free tier
• Domain & Hosting (if needed): $50/year
• Development Tools: Free (open source)
• Total Estimated Cost: $250

Technical Resources:
• GPU for training: 1x NVIDIA T4 (Google Colab)
• Development machines: Standard laptops
• Storage: 10GB for model and data
• Bandwidth: Standard internet connection"""
doc.add_paragraph(estimates_text)

doc.add_page_break()

# CONCLUSION
doc.add_heading('5. Conclusion', 1)
conclusion_text = """This project successfully demonstrates the application of modern NLP techniques to a low-resource indigenous language. The Kokborok POS tagging system represents a significant step toward digital inclusion for Kokborok-speaking communities.

Key Achievements:
• Developed first comprehensive POS tagger for Kokborok
• Created accessible web interface with educational content
• Achieved competitive accuracy using transfer learning
• Documented language history and cultural significance
• Established foundation for future Kokborok NLP tools

Impact:
The system enables linguists, educators, and native speakers to analyze Kokborok text automatically, supporting language preservation efforts and making the language more viable in the digital age.

Future Work:
• Expand to named entity recognition
• Develop machine translation capabilities
• Create mobile applications
• Build speech recognition systems
• Extend to related Tibeto-Burman languages"""
doc.add_paragraph(conclusion_text)

doc.add_page_break()

# REFERENCES
doc.add_heading('6. References', 1)
references = [
    'Conneau, A., et al. (2020). "Unsupervised Cross-lingual Representation Learning at Scale." ACL 2020.',
    'Devlin, J., et al. (2019). "BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding." NAACL 2019.',
    'Hugging Face Transformers Documentation. https://huggingface.co/docs/transformers',
    'Universal Dependencies Project. https://universaldependencies.org/',
    'Kokborok Language Resources. Tripura University.',
    'Gradio Documentation. https://gradio.app/docs',
    'PyTorch Documentation. https://pytorch.org/docs',
    'Bird, S., Klein, E., & Loper, E. (2009). "Natural Language Processing with Python." O\'Reilly Media.'
]

for i, ref in enumerate(references, 1):
    doc.add_paragraph(f'[{i}] {ref}', style='List Number')

# Save document
doc.save('Kokborok_NLP_Project_Report.docx')
print("Report generated successfully: Kokborok_NLP_Project_Report.docx")
