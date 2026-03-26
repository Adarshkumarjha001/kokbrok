from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a fresh document
doc = Document()

# Configure default style
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

print("Creating Title Page...")

# TITLE PAGE
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('KOKBOROK LANGUAGE NLP MODEL\n')
run.font.size = Pt(20)
run.font.bold = True
run.font.name = 'Times New Roman'

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Part-of-Speech Tagging System\n\n')
run.font.size = Pt(16)
run.font.name = 'Times New Roman'

project = doc.add_paragraph()
project.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = project.add_run('CAPSTONE PROJECT REPORT\n\n\n')
run.font.size = Pt(14)
run.font.bold = True
run.font.name = 'Times New Roman'

details = doc.add_paragraph()
details.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = details.add_run('Submitted by:\n')
run.font.size = Pt(12)
run = details.add_run('[Your Name]\n')
run.font.size = Pt(14)
run.font.bold = True
run = details.add_run('Roll No: [Your Roll Number]\n\n')
run.font.size = Pt(12)

dept = doc.add_paragraph()
dept.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = dept.add_run('Department of Computer Science & Engineering\n')
run.font.size = Pt(12)
run.font.bold = True

univ = doc.add_paragraph()
univ.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = univ.add_run('[University Name]\n')
run.font.size = Pt(12)
run.font.bold = True

year = doc.add_paragraph()
year.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = year.add_run('2024-2025')
run.font.size = Pt(12)

doc.add_page_break()

print("Creating Table of Contents...")

# TABLE OF CONTENTS
toc_heading = doc.add_paragraph()
toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = toc_heading.add_run('TABLE OF CONTENTS\n\n')
run.font.size = Pt(14)
run.font.bold = True

toc_items = [
    'Certificate .................................................. ii',
    'Acknowledgement .............................................. iii',
    'Abstract ..................................................... iv',
    'List of Figures .............................................. v',
    'List of Tables ............................................... vi',
    '',
    'CHAPTER 1: INTRODUCTION ...................................... 1',
    '  1.1 Objective of the Project ............................... 2',
    '  1.2 Description of the Project ............................. 3',
    '  1.3 Scope of the Project ................................... 5',
    '',
    'CHAPTER 2: SYSTEM DESCRIPTION ................................ 8',
    '  2.1 Customer/User Profiles ................................. 8',
    '  2.2 Assumptions and Dependencies ........................... 10',
    '  2.3 Functional Requirements ................................ 11',
    '  2.4 Non-Functional Requirements ............................ 13',
    '',
    'CHAPTER 3: DESIGN ............................................ 15',
    '  3.1 System Design .......................................... 15',
    '  3.2 Data Flow Diagrams ..................................... 17',
    '  3.3 Entity-Relationship Diagram ............................ 19',
    '  3.4 Database Design ........................................ 20',
    '',
    'CHAPTER 4: SCHEDULING AND ESTIMATES .......................... 22',
    '',
    'CHAPTER 5: CONCLUSION ........................................ 25',
    '',
    'REFERENCES ................................................... 27',
]

for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.left_indent = Inches(0.5) if item.startswith('  ') else Inches(0)

doc.add_page_break()

print("Saving document...")
doc.save('Kokborok_Clean_Report.docx')
print("Clean report created successfully: Kokborok_Clean_Report.docx")
