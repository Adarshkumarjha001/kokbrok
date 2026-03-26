from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Start with Part 3
doc = Document('Kokborok_Report_Part3.docx')

# 1.2 Description
doc.add_heading('1.2 Description of the Project', 2)
desc = """This project implements a comprehensive end-to-end Natural Language Processing solution specifically designed for the Kokborok language. The system encompasses three major components working in harmony to provide both functionality and education.

Component 1: Machine Learning Model
At the core of the system lies a sophisticated transformer-based neural network model built upon the XLM-RoBERTa (Cross-lingual Language Model - Robustly Optimized BERT Approach) architecture. This model has been fine-tuned specifically for token classification tasks in Kokborok language.

Model Specifications:
• Architecture: XLM-RoBERTa base model with 12 transformer layers
• Parameters: Approximately 125 million trainable parameters
• Hidden Dimensions: 768-dimensional vector representations
• Attention Mechanism: 12 attention heads per layer
• Vocabulary: 250,002 subword tokens using SentencePiece tokenization
• Output Classes: 17 distinct part-of-speech categories

The model leverages transfer learning, starting from a multilingual pre-trained model that has been exposed to 100+ languages during its initial training. This cross-lingual knowledge transfer is particularly valuable for Kokborok, which has limited digital resources and annotated corpora. The fine-tuning process adapts the model's general linguistic knowledge to the specific grammatical structures and patterns of Kokborok.

POS Categories Identified:
1. NOUN - Common nouns (e.g., objects, concepts)
2. PROPN - Proper nouns (names of people, places)
3. VERB - Action words and state verbs
4. ADJ - Adjectives describing nouns
5. ADV - Adverbs modifying verbs or adjectives
6. PRON - Pronouns replacing nouns
7. DET - Determiners (articles, quantifiers)
8. ADP - Adpositions (prepositions/postpositions)
9. AUX - Auxiliary verbs
10. CCONJ - Coordinating conjunctions
11. SCONJ - Subordinating conjunctions
12. PART - Particles
13. INTJ - Interjections
14. NUM - Numerals
15. PUNCT - Punctuation marks
16. X - Other/Foreign words
17. UNK - Unknown tokens

Component 2: Web Application
The system features a modern, responsive web application built using the Gradio framework, providing an intuitive interface for users of all technical levels.

Application Features:
• Real-time Text Analysis: Users can input Kokborok text and receive instant POS tagging results
• Visual Representation: Each word is color-coded according to its grammatical category
• Confidence Scores: The system displays prediction confidence for transparency
• Interactive Legend: A comprehensive guide explaining all POS tags
• Example Sentences: Pre-loaded examples for quick testing
• Multi-page Structure: Separate pages for different aspects of the project

The web interface is designed with accessibility in mind, featuring:
- Responsive design that works on desktop, tablet, and mobile devices
- High contrast colors meeting WCAG accessibility standards
- Clear typography and intuitive navigation
- Fast loading times and efficient resource usage

Component 3: Educational Content
Recognizing that technology alone cannot preserve a language, the system includes extensive educational materials:

History Section:
• Origins and etymology of Kokborok
• Historical timeline from ancient times to present
• Demographic information and speaker statistics
• Linguistic features and grammatical structures
• Cultural significance and traditional uses

Importance Section:
• Why indigenous language preservation matters
• The role of technology in language maintenance
• Real-world impact on communities
• Global context of language endangerment
• Future implications and opportunities

Technology Section:
• Detailed explanation of the NLP model
• How transformer architecture works
• The process of transfer learning
• System architecture and design decisions
• Technical specifications and performance metrics

Technical Implementation:
The system is implemented using modern software engineering practices:

Backend:
• Python 3.11+ for core logic
• PyTorch for deep learning operations
• Transformers library for model management
• Gradio for web framework and API
• Flask as alternative deployment option

Frontend:
• HTML5 for semantic structure
• CSS3 with Grid and Flexbox for responsive layout
• Vanilla JavaScript for interactivity
• No heavy frameworks for fast loading

Deployment:
• Hugging Face Spaces for model hosting
• GitHub Pages for static content (alternative)
• Cloud-based infrastructure for scalability
• Continuous deployment pipeline

The model weights are hosted on Hugging Face Model Hub, allowing the application to download them on-demand rather than bundling them with the application code. This approach reduces deployment size and allows for easy model updates without redeploying the entire application.
"""
doc.add_paragraph(desc).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.save('Kokborok_Report_Building.docx')
print("Building report...")
