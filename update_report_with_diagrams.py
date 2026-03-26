from docx import Document
from docx.shared import Inches

# Open existing report
doc = Document('Kokborok_NLP_Project_Report.docx')

# Find the Design section and add diagrams
# We'll insert after the "3.1.1 Data Flow Diagram" heading

# Add the diagrams
paragraphs = doc.paragraphs
for i, para in enumerate(paragraphs):
    if "3.1.1 Data Flow Diagram" in para.text or "Data Flow" in para.text:
        # Insert Level 0 DFD
        para_after = paragraphs[i+1]
        run = para_after.add_run()
        run.add_break()
        run.add_text("\n\nFigure 1: Level 0 DFD - Context Diagram\n")
        
        # Add image
        try:
            doc.add_picture('DFD_Level0_Context.png', width=Inches(6))
        except:
            print("Could not add Level 0 DFD image")
        
        doc.add_paragraph("\n")
        
        # Add Level 1 DFD
        doc.add_paragraph("Figure 2: Level 1 DFD - Detailed Process Flow\n")
        try:
            doc.add_picture('DFD_Level1_Detailed.png', width=Inches(6))
        except:
            print("Could not add Level 1 DFD image")
        
        break

# Find ER diagram section
for i, para in enumerate(paragraphs):
    if "3.1.1 E-R diagram" in para.text or "E-R" in para.text:
        para_after = paragraphs[i+1]
        run = para_after.add_run()
        run.add_break()
        run.add_text("\n\nFigure 3: Entity-Relationship Diagram\n")
        
        try:
            doc.add_picture('ER_Diagram.png', width=Inches(6))
        except:
            print("Could not add ER diagram image")
        
        break

# Save updated document
doc.save('Kokborok_NLP_Project_Report_with_Diagrams.docx')
print("Updated report saved as: Kokborok_NLP_Project_Report_with_Diagrams.docx")
