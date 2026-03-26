from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document('Kokborok_Clean_Report.docx')

print("Adding remaining sections...")

# 1.2 Description
doc.add_paragraph()
desc_heading = doc.add_paragraph()
run = desc_heading.add_run('1.2 Description of the Project')
run.font.size = Pt(13)
run.font.bold = True

desc_text = """This project implements a comprehensive NLP solution for Kokborok language consisting of three integrated components:

1. Machine Learning Model:
The core component is a transformer-based neural network built on XLM-RoBERTa architecture. The model specifications include:
• 125 million parameters
• 12 transformer layers with 768-dimensional hidden states
• 12 attention heads per layer
• 250,002 subword vocabulary using SentencePiece tokenization
• 17 output classes for POS categories

The model identifies: NOUN, PROPN, VERB, ADJ, ADV, PRON, DET, ADP, AUX, CCONJ, SCONJ, PART, INTJ, NUM, PUNCT, X, and UNK.

2. Web Application:
Built using Gradio framework, the application provides:
• Real-time text analysis with instant results
• Color-coded visualization of POS tags
• Confidence scores for each prediction
• Interactive legend explaining all tags
• Example sentences for testing
• Responsive design for all devices

3. Educational Content:
Comprehensive documentation including:
• Kokborok language history and origins
• Cultural significance and demographics
• Importance of language preservation
• Technical architecture and implementation details"""

p = doc.add_paragraph(desc_text)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# 1.3 Scope
doc.add_paragraph()
scope_heading = doc.add_paragraph()
run = scope_heading.add_run('1.3 Scope of the Project')
run.font.size = Pt(13)
run.font.bold = True

scope_text = """In Scope:
• Part-of-Speech tagging for 17 grammatical categories
• Web-based user interface with real-time analysis
• Support for Bengali and Roman scripts
• Educational content about Kokborok language
• Cloud deployment on Hugging Face Spaces
• API for programmatic access
• Confidence scoring for predictions

Out of Scope:
• Machine translation between languages
• Speech recognition or text-to-speech
• Named Entity Recognition (NER)
• Sentiment analysis
• Grammar correction tools
• Language learning applications

Future Enhancements:
• Extension to other Tibeto-Burman languages
• Integration with translation systems
• Mobile application development
• Voice interface capabilities
• Corpus building and annotation tools"""

p = doc.add_paragraph(scope_text)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.save('Kokborok_Complete_Report.docx')
print("Sections added successfully!")
