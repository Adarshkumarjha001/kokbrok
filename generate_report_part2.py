from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Open existing document
doc = Document('Kokborok_Morphological_Analysis_Full_Report.docx')

# CHAPTER 3: DESCRIPTION OF THE PROJECT
doc.add_page_break()
doc.add_heading('CHAPTER 3', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_heading('DESCRIPTION OF THE PROJECT', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('3.1 Project Overview', 2)
overview_text = """This project implements a comprehensive morphological analysis system for Kokborok using a multi-faceted approach that combines linguistic expertise, manual annotation, rule-based analysis, and state-of-the-art machine learning techniques.

<b>Core Components:</b>

The system consists of four primary components working in concert:

<b>1. Dataset Component:</b>
The foundation of the project is a carefully curated dataset of Kokborok words with detailed morphological annotations. This dataset includes:
• Surface forms (the actual words as they appear in text)
• Lemmas (dictionary forms)
• Root morphemes
• Affixes (prefixes, suffixes, infixes)
• Morphological features (number, person, tense, case, etc.)
• Part-of-speech tags
• Morpheme boundaries
• Gloss information

<b>2. Analysis Component:</b>
The morphological analyzer processes input words through multiple stages:
• Tokenization: Breaking text into individual words
• Normalization: Handling spelling variations and script differences
• Morpheme Segmentation: Identifying morpheme boundaries
• Feature Extraction: Determining grammatical features
• POS Classification: Assigning grammatical categories
• Confidence Scoring: Providing reliability measures

<b>3. Model Component:</b>
The machine learning component leverages XLM-RoBERTa, a multilingual transformer model:
• Pre-trained on 100+ languages including related Tibeto-Burman languages
• Fine-tuned on the Kokborok dataset
• 125 million parameters
• 12 transformer layers
• 768-dimensional hidden representations
• Capable of capturing complex morphological patterns

<b>4. Interface Component:</b>
A web-based interface provides access to the system:
• Interactive text input and analysis
• Real-time processing and visualization
• Educational content about Kokborok
• API endpoints for programmatic access
• Mobile-responsive design
• Multi-page documentation

<b>System Workflow:</b>

The complete analysis workflow proceeds as follows:

1. <b>Input Reception:</b> User provides Kokborok text through web interface or API
2. <b>Preprocessing:</b> Text is cleaned, normalized, and tokenized
3. <b>Initial Analysis:</b> Rule-based component attempts morpheme segmentation
4. <b>ML Processing:</b> Transformer model processes tokens and predicts POS tags
5. <b>Feature Extraction:</b> Grammatical features are identified based on morphemes
6. <b>Post-processing:</b> Results are formatted and validated
7. <b>Output Generation:</b> Analyzed text with annotations is returned to user
8. <b>Visualization:</b> Results are displayed with color-coding and explanations

<b>Technical Architecture:</b>

The system employs a three-tier architecture:

<b>Presentation Tier:</b>
• HTML5/CSS3/JavaScript frontend
• Responsive design using modern CSS
• Interactive visualizations
• Multiple pages (Home, History, Importance, Technology)

<b>Application Tier:</b>
• Python-based backend using Gradio framework
• RESTful API endpoints
• Request validation and error handling
• Session management
• Logging and monitoring

<b>Data