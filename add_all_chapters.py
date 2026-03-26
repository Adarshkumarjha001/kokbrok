from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document('Kokborok_Clean_Report.docx')

print("Adding Chapter 1: Introduction...")

# CHAPTER 1
doc.add_page_break()
ch1 = doc.add_paragraph()
ch1.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = ch1.add_run('CHAPTER 1\nINTRODUCTION\n\n')
run.font.size = Pt(14)
run.font.bold = True

intro = """Language serves as the cornerstone of cultural identity and heritage. For indigenous communities worldwide, their native languages embody centuries of accumulated knowledge, unique worldviews, and irreplaceable cultural traditions. Kokborok, a Tibeto-Burman language spoken by approximately one million people in Tripura, India, represents such a vital linguistic and cultural treasure.

Despite being recognized as an official language of Tripura since 1979, Kokborok faces significant challenges in the digital age. The lack of computational tools and resources creates a digital divide, forcing speakers to use dominant languages for online communication, education, and professional activities. This linguistic shift, particularly among younger generations, threatens the language's vitality and cultural continuity.

Natural Language Processing (NLP) offers a powerful solution to bridge this digital divide. By developing computational tools that understand and process Kokborok, we can enable speakers to use their native language in digital contexts, thereby supporting language maintenance and revitalization efforts.

This project addresses this critical need by developing a comprehensive Part-of-Speech (POS) tagging system for Kokborok. POS tagging, which identifies the grammatical category of each word in text, is a fundamental NLP task that serves as the foundation for more advanced applications such as machine translation, information extraction, and text generation."""

p = doc.add_paragraph(intro)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# 1.1 Objective
doc.add_paragraph()
obj_heading = doc.add_paragraph()
run = obj_heading.add_run('1.1 Objective of the Project')
run.font.size = Pt(13)
run.font.bold = True

obj_text = """The primary objectives of this project encompass technical, socio-cultural, and research dimensions:

Technical Objectives:
• Develop an accurate POS tagging model capable of identifying 17 different grammatical categories in Kokborok text
• Implement transfer learning using pre-trained XLM-RoBERTa transformer model
• Create a user-friendly web application for real-time text analysis
• Design scalable architecture deployable on cloud infrastructure
• Establish foundation for future Kokborok NLP applications

Socio-Cultural Objectives:
• Contribute to Kokborok language preservation through digital tools
• Educate users about language history and cultural significance
• Bridge the digital divide for Kokborok-speaking communities
• Empower linguists, educators, and native speakers
• Inspire similar efforts for other endangered languages

Research Objectives:
• Investigate transfer learning effectiveness for Tibeto-Burman languages
• Document best practices for low-resource language NLP
• Contribute to computational linguistics research"""

p = doc.add_paragraph(obj_text)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.save('Kokborok_Clean_Report.docx')
print("Chapter 1 added. Continuing...")
