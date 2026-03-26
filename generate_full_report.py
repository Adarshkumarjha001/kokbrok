from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_page_number(doc):
    """Add page numbers to footer"""
    for section in doc.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        
# Create document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# TITLE PAGE
title = doc.add_heading('Morphological Analysis of the Kokborok Language', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.size = Pt(18)
title.runs[0].font.bold = True

subtitle = doc.add_paragraph('Using a Custom Dataset for Part-of-Speech Tagging')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(16)

doc.add_paragraph('\n\n')
doc.add_paragraph('Capstone Project – I').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('CSE3391').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('\n\n\n')

# Project details
details = doc.add_paragraph()
details.alignment = WD_ALIGN_PARAGRAPH.CENTER
details.add_run('Submitted by:\n')
details.add_run('[Your Name]\n')
details.add_run('Roll No: [Your Roll Number]\n\n')
details.add_run('Under the Guidance of:\n')
details.add_run('[Guide Name]\n\n')
details.add_run('Department of Computer Science & Engineering\n')
details.add_run('[University Name]\n')
details.add_run('Academic Year: 2024-2025\n')

doc.add_page_break()

# CERTIFICATE PAGE
doc.add_heading('CERTIFICATE', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
cert_text = """This is to certify that the project entitled "Morphological Analysis of the Kokborok Language Using a Custom Dataset" submitted by [Your Name], Roll No. [Your Roll Number], in partial fulfillment of the requirements for the award of the degree of Bachelor of Technology in Computer Science and Engineering is a record of bonafide work carried out by him/her under my guidance and supervision.

The results embodied in this project have not been submitted to any other university or institution for the award of any degree or diploma.
"""
doc.add_paragraph(cert_text)
doc.add_paragraph('\n\n\n')
doc.add_paragraph('Guide Name: _________________')
doc.add_paragraph('Signature: _________________')
doc.add_paragraph('Date: _________________')

doc.add_page_break()

# ACKNOWLEDGEMENT
doc.add_heading('ACKNOWLEDGEMENT', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
ack_text = """I would like to express my sincere gratitude to all those who have contributed to the successful completion of this project.

First and foremost, I am deeply grateful to my project guide [Guide Name] for their invaluable guidance, continuous support, and encouragement throughout this project. Their expertise and insights have been instrumental in shaping this work.

I extend my heartfelt thanks to [HOD Name], Head of the Department of Computer Science and Engineering, for providing the necessary facilities and creating an environment conducive to research and learning.

I am thankful to the Kokborok-speaking community members who provided linguistic insights and helped validate the dataset. Special thanks to linguists and researchers working on Tibeto-Burman languages whose prior work formed the foundation of this study.

I acknowledge the support of my family and friends who encouraged me throughout this journey. Their patience and understanding have been a constant source of motivation.

Finally, I am grateful to all the faculty members and staff of the Computer Science and Engineering department for their support and cooperation.
"""
doc.add_paragraph(ack_text)

doc.add_page_break()

# ABSTRACT
doc.add_heading('ABSTRACT', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
abstract_text = """Kokborok, a Tibeto-Burman language spoken by approximately one million people in Northeast India, faces significant challenges in the digital age due to limited computational linguistic resources. This project addresses this gap by developing a comprehensive morphological analysis system for Kokborok, focusing on Part-of-Speech (POS) tagging using advanced Natural Language Processing techniques.

The project employs XLM-RoBERTa, a state-of-the-art multilingual transformer model, fine-tuned on a carefully curated Kokborok dataset. The system identifies and classifies 17 distinct grammatical categories including nouns, verbs, adjectives, pronouns, adverbs, determiners, adpositions, conjunctions, and punctuation marks. The morphological analyzer decomposes Kokborok words into their constituent morphemes, identifying roots, affixes, inflectional markers, and derivational patterns.

Key contributions of this work include: (1) Creation of an annotated Kokborok corpus with morphological tags, (2) Development of a rule-based morphological analyzer combined with machine learning approaches, (3) Implementation of a web-based interface for real-time text analysis, and (4) Comprehensive documentation of Kokborok morphological patterns.

The system achieves competitive accuracy on test data, demonstrating the effectiveness of transfer learning for low-resource languages. This work establishes a foundation for future NLP applications in Kokborok, including machine translation, information retrieval, and language learning tools. The project also contributes to language preservation efforts by creating digital resources for an indigenous language facing endangerment pressures.

Keywords: Kokborok, Morphological Analysis, Part-of-Speech Tagging, XLM-RoBERTa, Low-Resource Languages, Natural Language Processing, Tibeto-Burman Languages
"""
doc.add_paragraph(abstract_text)

doc.add_page_break()

# TABLE OF CONTENTS
doc.add_heading('TABLE OF CONTENTS', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
toc_items = [
    ('1. INTRODUCTION', '1'),
    ('   1.1 Purpose', '2'),
    ('   1.2 Scope', '3'),
    ('   1.3 Definitions and Terminology', '4'),
    ('   1.4 Background of Kokborok Language', '5'),
    ('2. OBJECTIVES OF THE PROJECT', '8'),
    ('   2.1 Primary Objectives', '8'),
    ('   2.2 Secondary Objectives', '9'),
    ('3. DESCRIPTION OF THE PROJECT', '10'),
    ('   3.1 Project Overview', '10'),
    ('   3.2 Methodology', '11'),
    ('   3.3 Technologies Used', '12'),
    ('4. LITERATURE REVIEW', '14'),
    ('   4.1 Morphological Analysis in NLP', '14'),
    ('   4.2 Work on Tibeto-Burman Languages', '16'),
    ('   4.3 Transfer Learning for Low-Resource Languages', '18'),
    ('5. SYSTEM DESCRIPTION', '20'),
    ('   5.1 Customer/User Profiles', '20'),
    ('   5.2 Assumptions and Dependencies', '21'),
    ('   5.3 Functional Requirements', '22'),
    ('   5.4 Non-Functional Requirements', '24'),
    ('6. MORPHOLOGICAL ANALYSIS OF KOKBOROK', '26'),
    ('   6.1 Phonological Structure', '26'),
    ('   6.2 Noun Morphology', '28'),
    ('   6.3 Verb Morphology', '31'),
    ('   6.4 Adjective and Adverb Morphology', '35'),
    ('   6.5 Case Markers and Postpositions', '37'),
    ('   6.6 Derivational Morphology', '39'),
    ('7. DATASET CONSTRUCTION', '42'),
    ('   7.1 Data Collection', '42'),
    ('   7.2 Annotation Process', '43'),
    ('   7.3 Dataset Statistics', '45'),
    ('8. SYSTEM DESIGN', '47'),
    ('   8.1 Architecture Overview', '47'),
    ('   8.2 Data Flow Diagrams', '49'),
    ('   8.3 Entity-Relationship Diagram', '51'),
    ('   8.4 Database Design', '52'),
    ('9. IMPLEMENTATION', '54'),
    ('   9.1 Model Training', '54'),
    ('   9.2 Web Interface Development', '56'),
    ('   9.3 API Development', '58'),
    ('10. TESTING AND EVALUATION', '60'),
    ('   10.1 Test Cases', '60'),
    ('   10.2 Performance Metrics', '62'),
    ('   10.3 Results and Analysis', '64'),
    ('11. SCHEDULING AND ESTIMATES', '66'),
    ('12. CONCLUSION AND FUTURE WORK', '68'),
    ('   12.1 Conclusion', '68'),
    ('   12.2 Limitations', '69'),
    ('   12.3 Future Enhancements', '70'),
    ('REFERENCES', '72'),
    ('APPENDICES', '74'),
]

for item, page in toc_items:
    p = doc.add_paragraph()
    p.add_run(item).bold = not item.startswith('   ')
    p.add_run('\t' + page)

doc.add_page_break()

# CHAPTER 1: INTRODUCTION
doc.add_heading('CHAPTER 1', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_heading('INTRODUCTION', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER

intro_text = """Morphology is a fundamental component of natural language processing that examines how words are formed, structured, and modified to convey meaning and grammatical relationships. For most Indian languages, particularly those with limited digital resources like Kokborok, the lack of annotated datasets and computational linguistic tools presents significant challenges for automated morphological analysis. However, such analysis is essential for developing comprehensive NLP applications and preserving linguistic heritage.

Kokborok, also known as Tripuri or Tiprakok, is the native language of the Tripuri people and serves as one of the official languages of Tripura state in Northeast India. As a member of the Tibeto-Burman language family, Kokborok exhibits rich morphological characteristics including agglutination, complex inflectional patterns, and extensive use of affixation. These linguistic features make it an excellent candidate for computational morphological investigation while simultaneously presenting unique challenges for NLP system development.

The significance of this project extends beyond mere technical achievement. Kokborok, like many indigenous languages worldwide, faces endangerment pressures from dominant languages. With approximately one million speakers, the language maintains vitality in its home region, but younger generations increasingly prefer Hindi, Bengali, or English for education and professional advancement. Digital language resources play a crucial role in language maintenance by demonstrating the language's relevance in modern technological contexts and providing tools for education, communication, and cultural preservation.

This project focuses on building a structured morphological analysis system for Kokborok using a carefully curated and annotated dataset. The system applies a combination of rule-based linguistic analysis, statistical methods, and state-of-the-art machine learning techniques to identify root words, affixes, inflectional markers, derivational patterns, and grammatical features. The implementation leverages XLM-RoBERTa, a multilingual transformer model, fine-tuned specifically for Kokborok Part-of-Speech tagging.

The morphological analyzer developed in this project can decompose Kokborok words into their constituent morphemes, identify grammatical categories, and provide detailed linguistic annotations. This capability forms the foundation for numerous downstream NLP applications including machine translation, information retrieval, text-to-speech systems, and automated language learning tools.
"""
doc.add_paragraph(intro_text)

# 1.1 PURPOSE
doc.add_heading('1.1 Purpose', 2)
purpose_text = """The primary purpose of this project is to develop a comprehensive morphological analysis system for the Kokborok language that can serve both academic research and practical NLP applications. Specific purposes include:

<b>Linguistic Documentation:</b> Create a systematic digital record of Kokborok morphological patterns, including inflectional and derivational processes. This documentation preserves linguistic knowledge that might otherwise remain undocumented or be lost as language use patterns change.

<b>Computational Resource Development:</b> Build machine-readable datasets and computational tools that enable automated processing of Kokborok text. These resources address the critical shortage of NLP tools for low-resource languages.

<b>Morpheme Identification:</b> Develop algorithms capable of identifying and classifying various morphological units including:
• Root words and stems
• Prefixes and suffixes
• Inflectional markers
• Derivational affixes
• Clitics and particles
• Compound word components

<b>Grammatical Feature Extraction:</b> Enable automatic identification of grammatical features such as:
• Number (singular, plural, dual)
• Person (first, second, third)
• Tense, aspect, and mood (TAM) markers
• Case markers
• Gender markers (where applicable)
• Definiteness and specificity
• Honorific markers

<b>POS Tagging:</b> Implement accurate Part-of-Speech tagging for Kokborok text, classifying words into 17 distinct grammatical categories based on their morphological and syntactic properties.

<b>Foundation for Advanced NLP:</b> Establish a baseline computational framework that future researchers and developers can build upon for more sophisticated NLP applications such as:
• Machine translation systems
• Information extraction tools
• Sentiment analysis
• Text summarization
• Question answering systems
• Language learning applications

<b>Language Preservation:</b> Contribute to efforts to preserve and promote Kokborok by demonstrating its compatibility with modern technology and creating resources that encourage its continued use in digital contexts.

<b>Academic Contribution:</b> Add to the body of research on Tibeto-Burman languages and low-resource language processing, providing insights that may be applicable to related languages facing similar challenges.
"""
doc.add_paragraph(purpose_text)

doc.add_page_break()

# 1.2 SCOPE
doc.add_heading('1.2 Scope', 2)
scope_text = """The scope of this project encompasses several interconnected components, each contributing to the overall goal of comprehensive morphological analysis for Kokborok.

<b>Within Scope:</b>

<b>Data Collection and Curation:</b>
• Gathering Kokborok lexical items from multiple sources including dictionaries, literary texts, newspapers, and community contributions
• Collecting approximately 1,000-2,000 annotated word forms representing diverse morphological patterns
• Ensuring representation of different word classes, inflectional paradigms, and derivational processes
• Documenting both Roman and Bengali script representations where applicable

<b>Morphological Annotation:</b>
• Manual annotation of morphological features by linguistically trained annotators
• Identification of morpheme boundaries within words
• Classification of morphemes by type (root, affix, clitic)
• Tagging of grammatical features (number, person, tense, case, etc.)
• Quality assurance through inter-annotator agreement checks

<b>Linguistic Analysis:</b>
• Detailed examination of Kokborok word formation processes
• Documentation of inflectional paradigms for major word classes
• Analysis of derivational morphology and productivity patterns
• Study of morphophonological alternations
• Comparison with related Tibeto-Burman languages

<b>Computational System Development:</b>
• Implementation of rule-based morphological analyzer using finite-state methods
• Training of machine learning models for POS tagging
• Development of morpheme segmentation algorithms
• Creation of morphological generation capabilities
• Integration of rule-based and statistical approaches

<b>Web Interface Development:</b>
• Design and implementation of user-friendly web interface
• Real-time morphological analysis functionality
• Visualization of morphological structure
• Educational content about Kokborok language and morphology
• API endpoints for programmatic access

<b>Evaluation and Testing:</b>
• Systematic evaluation using standard NLP metrics
• Testing on held-out data
• Error analysis and system refinement
• User testing with native speakers and linguists

<b>Documentation:</b>
• Comprehensive technical documentation
• Linguistic analysis reports
• User guides and tutorials
• API documentation

<b>Out of Scope:</b>

The following aspects, while related and potentially valuable, are explicitly excluded from the current project scope:

• <b>Syntactic Analysis:</b> Parsing of sentence structure, dependency relations, and phrase structure grammar
• <b>Semantic Analysis:</b> Deep semantic interpretation, word sense disambiguation, semantic role labeling
• <b>Machine Translation:</b> Full translation systems between Kokborok and other languages
• <b>Speech Processing:</b> Speech recognition, text-to-speech synthesis, prosody analysis
• <b>Discourse Analysis:</b> Analysis of text structure beyond the word level
• <b>Pragmatic Analysis:</b> Context-dependent meaning interpretation
• <b>Dialect Variation:</b> Comprehensive coverage of all Kokborok dialects (focus on standard variety)
• <b>Historical Linguistics:</b> Diachronic analysis of morphological changes over time

<b>Future Enhancements:</b>

While not within the current scope, the following enhancements are identified for future development:

• Extension to other Tibeto-Burman languages (Manipuri, Bodo, Garo)
• Integration with machine translation systems
• Development of mobile applications
• Incorporation of speech recognition capabilities
• Expansion to handle dialectal variation
• Creation of language learning applications
• Development of grammar checking tools
• Integration with corpus linguistics platforms
"""
doc.add_paragraph(scope_text)

doc.add_page_break()

# Continue with more sections...
# 1.3 DEFINITIONS
doc.add_heading('1.3 Definitions and Terminology', 2)

definitions_text = """This section provides definitions of key technical and linguistic terms used throughout this report.

<b>Morphology:</b> The branch of linguistics that studies the internal structure of words and how they are formed from smaller meaningful units called morphemes.

<b>Morpheme:</b> The smallest meaningful unit in a language. Morphemes cannot be divided into smaller meaningful parts. For example, in "unbreakable," there are three morphemes: "un-" (negation), "break" (root), and "-able" (capability).

<b>Root/Stem:</b> The core morpheme that carries the primary lexical meaning of a word. In Kokborok, "bai" (rice) is a root that can take various affixes.

<b>Affix:</b> A morpheme that attaches to a root or stem to modify its meaning or grammatical function. Affixes include:
• <b>Prefix:</b> Attached before the root (e.g., "un-" in "undo")
• <b>Suffix:</b> Attached after the root (e.g., "-ni" plural marker in Kokborok)
• <b>Infix:</b> Inserted within the root (rare in Kokborok)
• <b>Circumfix:</b> Surrounds the root (combination of prefix and suffix)

<b>Inflection:</b> Morphological process that modifies a word to express grammatical categories such as tense, number, person, case, or gender without changing the word's basic meaning or part of speech. Example: "go" → "goes" → "went" → "going"

<b>Derivation:</b> Morphological process that creates new words by adding affixes to roots, often changing the part of speech or significantly altering meaning. Example: "teach" (verb) → "teacher" (noun)

<b>Agglutination:</b> A morphological typology where words are formed by stringing together morphemes, each with a distinct and consistent meaning. Kokborok exhibits agglutinative characteristics.

<b>Part-of-Speech (POS):</b> A category of words that have similar grammatical properties. Common POS categories include noun, verb, adjective, adverb, pronoun, preposition, conjunction, and interjection.

<b>Lemma:</b> The canonical or dictionary form of a word. For verbs, typically the infinitive; for nouns, typically the singular form.

<b>Token:</b> An individual instance of a word or symbol in text. "The cat sat on the mat" contains 6 tokens.

<b>Tokenization:</b> The process of breaking text into individual tokens (words, punctuation marks, etc.).

<b>Corpus (plural: Corpora):</b> A large collection of texts used for linguistic analysis or machine learning training.

<b>Annotation:</b> The process of adding linguistic information (such as POS tags, morphological features) to text data.

<b>Finite-State Automaton (FSA):</b> A computational model consisting of states and transitions, used in morphological analysis to recognize and generate word forms.

<b>Transformer:</b> A deep learning architecture based on self-attention mechanisms, highly effective for NLP tasks.

<b>XLM-RoBERTa:</b> A multilingual transformer model pre-trained on text in 100 languages, used as the base model for this project.

<b>Fine-tuning:</b> The process of adapting a pre-trained model to a specific task or dataset by continuing training on task-specific data.

<b>Transfer Learning:</b> A machine learning technique where knowledge gained from one task is applied to a different but related task.

<b>Low-Resource Language:</b> A language with limited digital resources, annotated corpora, and NLP tools available for computational processing.

<b>Tibeto-Burman:</b> A language family within Sino-Tibetan, including languages spoken across South Asia, Southeast Asia, and East Asia.

<b>Case Marker:</b> A morpheme that indicates the grammatical function of a noun phrase (subject, object, possessive, etc.).

<b>TAM (Tense-Aspect-Mood):</b> Grammatical categories that express when an action occurs (tense), how it unfolds over time (aspect), and the speaker's attitude toward it (mood).

<b>Clitic:</b> A morpheme that has syntactic characteristics of a word but is phonologically dependent on another word.

<b>Morphophonology:</b> The study of how morphological and phonological processes interact, including sound changes that occur at morpheme boundaries.

<b>Productivity:</b> The degree to which a morphological process can be used to create new words. Highly productive processes can be applied to many roots.

<b>Allomorph:</b> Different phonological forms of the same morpheme that occur in different contexts. Example: English plural "-s" vs. "-es" vs. "-en"

<b>Compounding:</b> A word formation process that combines two or more independent words to create a new word with a distinct meaning.
"""
doc.add_paragraph(definitions_text)

doc.add_page_break()

# 1.4 BACKGROUND OF KOKBOROK LANGUAGE
doc.add_heading('1.4 Background of Kokborok Language', 2)
background_text = """<b>1.4.1 Historical and Cultural Context</b>

Kokborok, the ancestral language of the Tripuri people, has a rich history spanning centuries in the northeastern region of India. The name "Kokborok" derives from "Kok" (language/verbal) and "Borok" (people), literally meaning "the language of the Borok people." The language serves as a crucial marker of ethnic identity for the Tripuri community and plays a central role in preserving their cultural heritage.

Historically, the Tripuri kingdom (also known as the Manikya dynasty) ruled the region from approximately the 14th century until 1949, when it merged with India. During this period, Kokborok flourished as the language of the royal court, administration, and common people. However, the influence of Bengali culture and language, particularly during the later period of the kingdom and after integration with India, led to Bengali becoming the dominant language of administration and education.

<b>1.4.2 Linguistic Classification</b>

Kokborok belongs to the Tibeto-Burman branch of the Sino-Tibetan language family. Within Tibeto-Burman, it is classified under the Bodo-Garo group, which includes related languages such as:
• Bodo (spoken in Assam)
• Garo (spoken in Meghalaya)
• Dimasa (spoken in Assam)
• Rabha (spoken in Assam and West Bengal)

This linguistic affiliation connects Kokborok to a broader family of languages spoken across South Asia, Southeast Asia, and East Asia, sharing certain typological features and historical origins.

<b>1.4.3 Geographic Distribution</b>

Kokborok is primarily spoken in:
• <b>Tripura, India:</b> The main concentration of speakers, where it is an official language
• <b>Assam, India:</b> Significant Tripuri populations in districts bordering Tripura
• <b>Mizoram, India:</b> Smaller communities in northern regions
• <b>Bangladesh:</b> Tripuri communities in Chittagong Hill Tracts and Sylhet division
• <b>Diaspora Communities:</b> Tripuri populations in other parts of India and abroad

The total number of speakers is estimated at approximately 1 million, though exact figures vary depending on the source and methodology used for counting.

<b>1.4.4 Dialectal Variation</b>

Kokborok exhibits dialectal variation across its geographic range. Major dialects include:
• <b>Debbarma:</b> The most widely spoken dialect, considered the standard variety
• <b>Reang (Bru):</b> Spoken by the Reang community
• <b>Jamatia:</b> Spoken by the Jamatia tribe
• <b>Noatia:</b> Spoken by the Noatia community
• <b>Koloi:</b> Spoken in certain regions

While these dialects share core vocabulary and grammatical structures, they differ in phonology, lexicon, and some morphological patterns. This project focuses primarily on the Debbarma dialect, which serves as the basis for standardization efforts.

<b>1.4.5 Writing Systems</b>

Kokborok has been written in multiple scripts throughout its history:

<b>Bengali Script:</b> Historically, Kokborok was written using the Bengali script (Bangla lipi), which was the administrative script of the region. Many older texts, including religious and literary works, use this script.

<b>Roman Script:</b> In 1979, the Roman script was officially adopted for Kokborok to improve accessibility and ease of learning. The Roman script system was developed by the Kokborok Sahitya Sabha (Kokborok Literary Society) and has since become the primary writing system for education and official purposes.

<b>Koloma Script:</b> An indigenous script historically used by the Tripuri people, though it fell out of common use. Recent revival efforts aim to preserve this cultural heritage.

The choice of script remains somewhat contentious, with different communities and organizations preferring different systems. This project uses Roman script representation, which is most widely used in digital contexts.

<b>1.4.6 Official Status and Recognition</b>

Kokborok achieved significant milestones in official recognition:
• <b>1979:</b> Recognized as an official language of Tripura state
• <b>1979:</b> Roman script officially adopted
• <b>2003:</b> Included in the school curriculum by Tripura Board of Secondary Education
• <b>Present:</b> Used in state government communications, education, media, and literature

Despite official recognition, Kokborok faces challenges in practical implementation, with limited resources for education, publishing, and digital technology.

<b>1.4.7 Sociolinguistic Situation</b>

The sociolinguistic situation of Kokborok is complex:

<b>Language Vitality:</b> While Kokborok maintains strong vitality in rural areas and among older generations, younger urban speakers increasingly prefer Hindi, Bengali, or English for education and professional advancement.

<b>Multilingualism:</b> Most Kokborok speakers are multilingual, typically speaking Bengali and/or Hindi in addition to their native language. This multilingualism, while enriching, also creates pressure toward language shift.

<b>Domain Restriction:</b> Kokborok is primarily used in home and community settings, with limited presence in higher education, technology, business, and formal domains.

<b>Intergenerational Transmission:</b> While generally maintained in rural areas, urban families show declining rates of language transmission to children.

<b>1.4.8 Linguistic Characteristics</b>

Kokborok exhibits several distinctive linguistic features:

<b>Phonology:</b>
• Relatively simple consonant inventory
• Six vowel phonemes
• Tone is not phonemic (unlike many Sino-Tibetan languages)
• Syllable structure predominantly CV (consonant-vowel)

<b>Morphology:</b>
• Agglutinative tendencies with extensive use of suffixes
• Rich system of verbal inflection for tense, aspect, mood
• Noun classification and numeral classifiers
• Productive derivational processes

<b>Syntax:</b>
• Subject-Object-Verb (SOV) word order
• Postpositions rather than prepositions
• Extensive use of particles for grammatical functions
• Relative clauses precede head nouns

<b>Lexicon:</b>
• Core vocabulary of Tibeto-Burman origin
• Significant borrowings from Bengali, Sanskrit, and English
• Specialized vocabulary for traditional culture, agriculture, and social organization

<b>1.4.9 Current Challenges</b>

Kokborok faces several challenges in the modern era:

<b>Digital Divide:</b> Limited digital resources, including keyboards, fonts, spell-checkers, and online content
<b>Educational Resources:</b> Shortage of quality textbooks, teaching materials, and trained teachers
<b>Standardization:</b> Ongoing debates about script choice, spelling conventions, and dialect standardization
<b>Language Shift:</b> Pressure from dominant languages in education, media, and employment
<b>Documentation:</b> Insufficient linguistic documentation and description
<b>Technology Gap:</b> Lack of NLP tools, translation systems, and language technology

This project directly addresses the technology gap by creating computational resources for Kokborok morphological analysis.
"""
doc.add_paragraph(background_text)

doc.add_page_break()

# CHAPTER 2: OBJECTIVES
doc.add_heading('CHAPTER 2', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_heading('OBJECTIVES OF THE PROJECT', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('2.1 Primary Objectives', 2)
primary_obj = """The primary objectives of this project are carefully designed to address the critical gap in computational linguistic resources for Kokborok while contributing to both theoretical understanding and practical applications.

<b>Objective 1: Develop a Comprehensive Morphological Analyzer</b>

Create a robust computational system capable of analyzing Kokborok words at the morphological level. This analyzer should:
• Decompose words into constituent morphemes (roots, affixes, clitics)
• Identify morpheme boundaries accurately
• Classify morphemes by type and function
• Handle both inflectional and derivational morphology
• Process words in real-time with high accuracy
• Provide detailed morphological annotations

Success Criteria: The analyzer should achieve at least 85% accuracy on a held-out test set of annotated Kokborok words.

<b>Objective 2: Construct a High-Quality Annotated Dataset</b>

Build a linguistically accurate, machine-readable dataset of Kokborok words with comprehensive morphological annotations. The dataset should:
• Contain 1,000-2,000 unique word forms
• Cover diverse morphological patterns and word classes
• Include detailed annotations for each word
• Maintain high inter-annotator agreement (>90%)
• Be formatted for easy use in NLP applications
• Be publicly available for research purposes

Success Criteria: Dataset should be validated by native speakers and linguists, with documented annotation guidelines and quality metrics.

<b>Objective 3: Implement Accurate Part-of-Speech Tagging</b>

Develop a POS tagging system specifically tuned for Kokborok that can:
• Classify words into 17 distinct grammatical categories
• Achieve state-of-the-art accuracy for the language
• Handle ambiguous cases appropriately
• Provide confidence scores for predictions
• Process text efficiently in real-time
• Integrate with the morphological analyzer

Success Criteria: POS tagging accuracy should exceed 90% on test data, comparable to or better than existing systems for related languages.

<b>Objective 4: Extract and Document Morphological Patterns</b>

Conduct systematic linguistic analysis to identify and document:
• Inflectional paradigms for major word classes
• Derivational processes and their productivity
• Morphophonological rules and alternations
• Compound word formation patterns
• Grammatical feature systems (number, person, tense, case, etc.)
• Morphological typology of Kokborok

Success Criteria: Comprehensive documentation of morphological patterns validated by linguistic experts and native speakers.

<b>Objective 5: Create an Accessible Web-Based Interface</b>

Develop a user-friendly web application that:
• Allows real-time morphological analysis of Kokborok text
• Visualizes morphological structure clearly
• Provides educational content about the language
• Offers API access for developers
• Works across devices (desktop, tablet, mobile)
• Requires no technical expertise to use

Success Criteria: Positive user feedback from at least 80% of test users, including native speakers, linguists, and students.
"""
doc.add_paragraph(primary_obj)

doc.add_page_break()

doc.add_heading('2.2 Secondary Objectives', 2)
secondary_obj = """Beyond the primary goals, this project pursues several secondary objectives that enhance its value and impact:

<b>Objective 6: Contribute to Language Preservation</b>

Support efforts to preserve and promote Kokborok by:
• Creating digital resources that demonstrate the language's relevance in technology
• Documenting linguistic patterns that might otherwise be lost
• Providing tools that encourage language use among younger generations
• Making the language more accessible to learners and researchers
• Raising awareness about the importance of indigenous language preservation

<b>Objective 7: Establish Foundation for Advanced NLP Applications</b>

Create a baseline system that enables future development of:
• Machine translation systems (Kokborok ↔ English/Bengali/Hindi)
• Information retrieval and search engines for Kokborok content
• Text-to-speech and speech recognition systems
• Grammar checking and writing assistance tools
• Language learning applications
• Corpus linguistics research platforms

<b>Objective 8: Advance Low-Resource Language Processing</b>

Contribute to the broader field of NLP for low-resource languages by:
• Demonstrating effective transfer learning techniques
• Developing methodologies applicable to other Tibeto-Burman languages
• Creating reusable tools and frameworks
• Publishing findings and resources for the research community
• Establishing best practices for dataset creation and annotation

<b>Objective 9: Foster Interdisciplinary Collaboration</b>

Bridge the gap between linguistics and computer science by:
• Combining linguistic expertise with computational methods
• Creating resources useful for both linguists and NLP researchers
• Demonstrating the value of interdisciplinary approaches
• Encouraging collaboration between academic institutions and language communities

<b>Objective 10: Promote Digital Inclusion</b>

Support digital inclusion for Kokborok speakers by:
• Enabling use of the language in digital contexts
• Reducing barriers to technology access
• Demonstrating that indigenous languages can thrive in the digital age
• Inspiring similar efforts for other marginalized languages

<b>Objective 11: Educational Impact</b>

Create educational value through:
• Resources for teaching Kokborok grammar
• Tools for language learning and practice
• Documentation accessible to students and educators
• Demonstrations of NLP concepts using a real-world language
• Inspiration for students from the Tripuri community to pursue technology careers

<b>Objective 12: Research Contribution</b>

Add to academic knowledge by:
• Publishing findings in conferences and journals
• Making datasets and code publicly available
• Documenting methodologies for replication
• Providing baseline results for future comparison
• Contributing to understanding of Tibeto-Burman morphology
"""
doc.add_paragraph(secondary_obj)

doc.add_page_break()

# Save the document
doc.save('Kokborok_Morphological_Analysis_Full_Report.docx')
print("Full report generated successfully!")
print("File: Kokborok_Morphological_Analysis_Full_Report.docx")
print("This is Part 1 of the report (approximately 15-20 pages)")
print("Run the continuation script to add remaining chapters...")
