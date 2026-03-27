from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Load the first part
doc = Document('Kokborok_Report_Part1.docx')

# ==================== TABLE OF CONTENTS ====================
doc.add_heading('TABLE OF CONTENTS', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER

toc_data = [
    ('Certificate', 'ii'),
    ('Acknowledgement', 'iii'),
    ('Abstract', 'iv'),
    ('Table of Contents', 'v'),
    ('List of Figures', 'vi'),
    ('List of Tables', 'vii'),
    ('', ''),
    ('CHAPTER 1: INTRODUCTION', '1'),
    ('1.1 Objective of the Project', '2'),
    ('1.2 Description of the Project', '3'),
    ('1.3 Scope of the Project', '5'),
    ('1.3.1 Use Case Model', '6'),
    ('1.4 Organization of Report', '7'),
    ('', ''),
    ('CHAPTER 2: LITERATURE REVIEW', '8'),
    ('2.1 Overview of NLP for Low-Resource Languages', '8'),
    ('2.2 Transformer Models and Transfer Learning', '10'),
    ('2.3 Part-of-Speech Tagging Techniques', '12'),
    ('2.4 Related Work', '14'),
    ('', ''),
    ('CHAPTER 3: SYSTEM DESCRIPTION', '16'),
    ('3.1 Customer/User Profiles', '16'),
    ('3.2 Assumptions and Dependencies', '18'),
    ('3.3 Functional Requirements', '19'),
    ('3.4 Non-Functional Requirements', '21'),
    ('', ''),
    ('CHAPTER 4: SYSTEM DESIGN', '23'),
    ('4.1 System Architecture', '23'),
    ('4.2 Data Flow Diagrams', '25'),
    ('4.2.1 Level 0 DFD (Context Diagram)', '25'),
    ('4.2.2 Level 1 DFD (Detailed Process)', '26'),
    ('4.3 Entity-Relationship Diagram', '28'),
    ('4.4 Database Design', '29'),
    ('4.5 User Interface Design', '31'),
    ('', ''),
    ('CHAPTER 5: IMPLEMENTATION', '33'),
    ('5.1 Technology Stack', '33'),
    ('5.2 Model Development', '35'),
    ('5.3 Web Application Development', '37'),
    ('5.4 Integration and Testing', '39'),
    ('', ''),
    ('CHAPTER 6: RESULTS AND DISCUSSION', '41'),
    ('6.1 Model Performance', '41'),
    ('6.2 System Evaluation', '43'),
    ('6.3 User Feedback', '45'),
    ('', ''),
    ('CHAPTER 7: CONCLUSION AND FUTURE WORK', '47'),
    ('7.1 Conclusion', '47'),
    ('7.2 Limitations', '48'),
    ('7.3 Future Enhancements', '49'),
    ('', ''),
    ('REFERENCES', '51'),
    ('APPENDICES', '53'),
]

for item, page in toc_data:
    if item == '':
        doc.add_paragraph()
    else:
        p = doc.add_paragraph()
        p.add_run(item).font.size = Pt(12)
        p.add_run('\t' * 5 + page).font.size = Pt(12)

doc.add_page_break()

# ==================== LIST OF FIGURES ====================
doc.add_heading('LIST OF FIGURES', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER

figures = [
    ('Figure 1.1', 'Use Case Diagram', '6'),
    ('Figure 4.1', 'System Architecture Diagram', '24'),
    ('Figure 4.2', 'Level 0 DFD - Context Diagram', '25'),
    ('Figure 4.3', 'Level 1 DFD - Detailed Process Flow', '27'),
    ('Figure 4.4', 'Entity-Relationship Diagram', '28'),
    ('Figure 4.5', 'Database Schema', '30'),
    ('Figure 4.6', 'User Interface Mockup - Home Page', '31'),
    ('Figure 4.7', 'User Interface Mockup - Analysis Page', '32'),
    ('Figure 5.1', 'Model Training Pipeline', '36'),
    ('Figure 5.2', 'Web Application Architecture', '38'),
    ('Figure 6.1', 'Model Accuracy Graph', '42'),
    ('Figure 6.2', 'Performance Metrics Comparison', '44'),
]

for fig_num, fig_name, page in figures:
    p = doc.add_paragraph()
    p.add_run(f'{fig_num}: {fig_name}').font.size = Pt(12)
    p.add_run('\t' * 3 + page).font.size = Pt(12)

doc.add_page_break()

# ==================== LIST OF TABLES ====================
doc.add_heading('LIST OF TABLES', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER

tables = [
    ('Table 2.1', 'Comparison of POS Tagging Approaches', '13'),
    ('Table 3.1', 'User Profile Summary', '17'),
    ('Table 3.2', 'Functional Requirements', '20'),
    ('Table 3.3', 'Non-Functional Requirements', '22'),
    ('Table 4.1', 'Database Tables and Attributes', '30'),
    ('Table 5.1', 'Technology Stack Components', '34'),
    ('Table 5.2', 'Model Hyperparameters', '36'),
    ('Table 6.1', 'POS Tagging Accuracy by Category', '42'),
    ('Table 6.2', 'System Performance Metrics', '44'),
    ('Table 7.1', 'Project Timeline and Milestones', '50'),
]

for table_num, table_name, page in tables:
    p = doc.add_paragraph()
    p.add_run(f'{table_num}: {table_name}').font.size = Pt(12)
    p.add_run('\t' * 3 + page).font.size = Pt(12)

doc.add_page_break()

doc.save('Kokborok_Report_Part2.docx')
print("Part 2 created...")
