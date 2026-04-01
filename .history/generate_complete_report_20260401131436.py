from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

# Create document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# ==================== TITLE PAGE ====================
title = doc.add_heading('KOKBOROK LANGUAGE NLP MODEL', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.size = Pt(24)
title.runs[0].font.bold = True

subtitle = doc.add_paragraph('Part-of-Speech Tagging System for Indigenous Language Preservation')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(16)
subtitle.runs[0].font.italic = True

doc.add_paragraph('\n\n')

# Project type
project_type = doc.add_paragraph('CAPSTONE PROJECT REPORT')
project_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
project_type.runs[0].font.size = Pt(14)
project_type.runs[0].font.bold = True

doc.add_paragraph('\n')

# Submitted info
submitted = doc.add_paragraph('Submitted in partial fulfillment of the requirements')
submitted.alignment = WD_ALIGN_PARAGRAPH.CENTER
submitted.runs[0].font.size = Pt(12)

submitted2 = doc.add_paragraph('for the degree of')
submitted2.alignment = WD_ALIGN_PARAGRAPH.CENTER
submitted2.runs[0].font.size = Pt(12)

degree = doc.add_paragraph('Bachelor of Technology')
degree.alignment = WD_ALIGN_PARAGRAPH.CENTER
degree.runs[0].font.size = Pt(14)
degree.runs[0].font.bold = True

doc.add_paragraph('\n\n')

# Student details
details = doc.add_paragraph()
details.alignment = WD_ALIGN_PARAGRAPH.CENTER
details.add_run('Submitted by:\n').font.size = Pt(12)
details.add_run('[Your Name]\n').font.size = Pt(14)
details.runs[-1].font.bold = True
details.add_run('Roll No: [Your Roll Number]\n').font.size = Pt(12)
details.add_run('Registration No: [Your Registration Number]\n\n').font.size = Pt(12)

details.add_run('Under the guidance of:\n').font.size = Pt(12)
details.add_run('[Guide Name]\n').font.size = Pt(14)
details.runs[-1].font.bold = True
details.add_run('[Designation]\n\n').font.size = Pt(12)

# Department
dept = doc.add_paragraph('Department of Computer Science & Engineering')
dept.alignment = WD_ALIGN_PARAGRAPH.CENTER
dept.runs[0].font.size = Pt(13)
dept.runs[0].font.bold = True

university = doc.add_paragraph('[University Name]')
university.alignment = WD_ALIGN_PARAGRAPH.CENTER
university.runs[0].font.size = Pt(13)
university.runs[0].font.bold = True

year = doc.add_paragraph(f'{datetime.now().year}')
year.alignment = WD_ALIGN_PARAGRAPH.CENTER
year.runs[0].font.size = Pt(13)

doc.add_page_break()

# ==================== CERTIFICATE ====================
doc.add_heading('CERTIFICATE', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER

cert_text = """This is to certify that the project entitled "KOKBOROK LANGUAGE NLP MODEL: Part-of-Speech Tagging System" submitted by [Your Name], Roll No: [Your Roll Number] in partial fulfillment of the requirements for the award of Bachelor of Technology in Computer Science & Engineering to [University Name] is a bonafide record of the work carried out by him/her under my supervision and guidance.

The project embodies results of original work, and studies are carried out by the student himself/herself and the contents of the project do not form the basis for the award of any other degree to the candidate or to anybody else from this or any other University/Institution.
"""
cert_para = doc.add_paragraph(cert_text)
cert_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_paragraph('\n\n')

# Signatures
sig_table = doc.add_table(rows=2, cols=2)
sig_table.cell(0, 0).text = '\n\n\n[Guide Name]\nProject Guide'
sig_table.cell(0, 1).text = '\n\n\n[HOD Name]\nHead of Department'
sig_table.cell(1, 0).text = 'Date: __________'
sig_table.cell(1, 1).text = 'Date: __________'

doc.add_page_break()

# ==================== ACKNOWLEDGEMENT ====================
doc.add_heading('ACKNOWLEDGEMENT', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER

ack_text = """I would like to express my sincere gratitude to all those who have contributed to the successful completion of this project.

First and foremost, I am deeply grateful to my project guide, [Guide Name], for their invaluable guidance, continuous support, and encouragement throughout this project. Their expertise and insights have been instrumental in shaping this work.

I extend my heartfelt thanks to [HOD Name], Head of the Department of Computer Science & Engineering, for providing the necessary facilities and creating an environment conducive to research and development.

I am thankful to all the faculty members of the Computer Science & Engineering department for their support and valuable suggestions during various stages of this project.

I would like to acknowledge the Kokborok-speaking community and linguists who provided valuable insights into the language structure and cultural context, which were essential for developing this NLP system.

I am grateful to Hugging Face for providing the pre-trained models and infrastructure that made this project possible, and to the open-source community for their excellent tools and libraries.

Finally, I express my gratitude to my family and friends for their constant support and encouragement throughout this endeavor.
"""
ack_para = doc.add_paragraph(ack_text)
ack_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_paragraph('\n\n')
signature = doc.add_paragraph('[Your Name]')
signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT

doc.add_page_break()

# ==================== ABSTRACT ====================
doc.add_heading('ABSTRACT', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER

abstract_text = """Language is the cornerstone of cultural identity, and for indigenous communities, preserving their native languages is crucial for maintaining their heritage. Kokborok, a Tibeto-Burman language spoken by approximately one million people in Tripura, India, faces challenges in the digital age due to limited technological resources and tools.

This project presents a comprehensive Natural Language Processing (NLP) system specifically designed for the Kokborok language, focusing on Part-of-Speech (POS) tagging. The system leverages state-of-the-art transformer architecture, specifically XLM-RoBERTa, fine-tuned for token classification to identify 17 different grammatical categories in Kokborok text.

The project addresses the critical need for digital language tools for low-resource indigenous languages. By utilizing transfer learning from multilingual pre-trained models, we overcome the challenge of limited annotated data availability. The system achieves competitive accuracy in POS tagging, demonstrating the effectiveness of modern NLP techniques for under-resourced languages.

The implementation includes a user-friendly web application built with Gradio framework, providing an interactive interface for real-time text analysis. The application features comprehensive educational content about Kokborok language history, cultural significance, and the importance of indigenous language preservation in the digital era.

The system architecture follows a modular design with three main components: a presentation layer (web interface), an application layer (API and processing logic), and a model layer (transformer-based POS tagger). The deployment strategy utilizes cloud infrastructure (Hugging Face Spaces) for accessibility and scalability.

This project contributes to the field of computational linguistics by providing the first comprehensive POS tagging system for Kokborok, establishing a foundation for future NLP applications such as machine translation, named entity recognition, and speech processing. The work demonstrates that with appropriate transfer learning techniques, effective NLP tools can be developed even for languages with limited digital resources.

Keywords: Natural Language Processing, Kokborok Language, Part-of-Speech Tagging, XLM-RoBERTa, Indigenous Languages, Transfer Learning, Low-Resource Languages, Language Preservation
"""
abstract_para = doc.add_paragraph(abstract_text)
abstract_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_page_break()

# Save first part
doc.save('Kokborok_Report_Part1.docxx')
print("Part 1 created...")
