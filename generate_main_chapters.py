from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document('Kokborok_Report_Part2.docx')

# ==================== CHAPTER 1: INTRODUCTION ====================
doc.add_page_break()
heading = doc.add_heading('CHAPTER 1', 1)
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_heading('INTRODUCTION', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER

intro_text = """Language is not merely a tool for communication; it is the repository of a community's history, culture, knowledge, and identity. For indigenous communities worldwide, their native languages represent centuries of accumulated wisdom, unique perspectives on the world, and irreplaceable cultural heritage. However, in our increasingly digital world, languages without technological support face an existential threat.

Kokborok, also known as Tripuri or Tiprakok, is a Tibeto-Burman language spoken by approximately one million people, primarily in the Indian state of Tripura and neighboring regions of Bangladesh and Myanmar. As one of the 22 scheduled languages of India and an official language of Tripura since 1979, Kokborok holds significant cultural and administrative importance. Despite this official recognition, the language faces challenges in the digital domain due to limited technological resources, tools, and applications.

The digital divide affecting indigenous languages like Kokborok is not merely a technological issue—it is a matter of cultural survival. When a language lacks digital tools, its speakers are compelled to use dominant languages for online communication, education, and professional work. This linguistic shift, particularly among younger generations, accelerates language endangerment and cultural erosion.

Natural Language Processing (NLP) offers a powerful solution to bridge this digital divide. By developing computational tools that understand and process indigenous languages, we can enable their speakers to use their native languages in digital spaces, thereby promoting language maintenance and revitalization. Part-of-Speech (POS) tagging, the process of marking words in text with their grammatical categories, is a fundamental NLP task that serves as the foundation for more advanced applications such as machine translation, information extraction, and text generation.

This project addresses the critical need for NLP tools for Kokborok by developing a comprehensive POS tagging system. Leveraging state-of-the-art transformer architecture and transfer learning techniques, we have created a system capable of accurately identifying 17 different parts of speech in Kokborok text. The system is accompanied by an accessible web application that not only provides the tagging functionality but also educates users about the language's history, cultural significance, and the importance of indigenous language preservation.
"""
doc.add_paragraph(intro_text).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# 1.1 Objective
doc.add_heading('1.1 Objective of the Project', 2)

obj_text = """The primary objectives of this project are multifaceted, addressing both technical and socio-cultural dimensions:

Technical Objectives:

1. Develop an Accurate POS Tagging Model: Create a robust machine learning model capable of identifying 17 different parts of speech in Kokborok text with high accuracy, despite the language being low-resource in terms of digital corpora and annotated datasets.

2. Implement Transfer Learning: Utilize pre-trained multilingual transformer models (XLM-RoBERTa) and fine-tune them for Kokborok, demonstrating the effectiveness of transfer learning for low-resource languages.

3. Create an Accessible Web Interface: Develop a user-friendly web application that allows users to analyze Kokborok text in real-time, with intuitive visualization of results through color-coded tags and confidence scores.

4. Ensure Scalability and Deployment: Design the system architecture to be scalable and deploy it on cloud infrastructure (Hugging Face Spaces) for global accessibility.

5. Establish a Foundation for Future NLP Tools: Create a baseline system that can serve as the foundation for more advanced Kokborok NLP applications such as named entity recognition, machine translation, and sentiment analysis.

Socio-Cultural Objectives:

1. Promote Language Preservation: Contribute to the preservation and promotion of Kokborok language by providing digital tools that make the language more viable in the modern technological landscape.

2. Educate and Raise Awareness: Develop comprehensive educational content about Kokborok language history, linguistic features, and cultural significance, raising awareness about the importance of indigenous language preservation.

3. Bridge the Digital Divide: Enable Kokborok speakers to use their native language in digital contexts, reducing the linguistic barrier that forces them to use dominant languages online.

4. Empower the Community: Provide tools that empower linguists, educators, students, and native speakers to work with their language in computational contexts.

5. Inspire Similar Efforts: Demonstrate that effective NLP tools can be developed for low-resource indigenous languages, potentially inspiring similar projects for other endangered languages.

Research Objectives:

1. Investigate Transfer Learning Effectiveness: Examine how well multilingual pre-trained models transfer to Tibeto-Burman languages with limited training data.

2. Document Best Practices: Establish and document best practices for developing NLP systems for low-resource indigenous languages.

3. Contribute to Computational Linguistics: Add to the body of knowledge in computational linguistics, particularly in the area of morphologically rich and agglutinative languages.
"""
doc.add_paragraph(obj_text).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.save('Kokborok_Report_Part3.docx')
print("Part 3 created...")
